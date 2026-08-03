"""
reportgen.py
Generates college-format project reports as PDF (ReportLab) and
DOCX (python-docx) from a data dictionary collected in Create Report.
"""

# pylint: disable=no-member
# ReportLab's BaseDocTemplate/Frame classes set attributes like leftMargin,
# rightMargin, width, height, page dynamically at runtime (inside __init__),
# not as static class attributes. Pylint's static analysis can't see these,
# so it raises false "no-member" errors throughout this file. This is a
# known, well-documented limitation with ReportLab specifically - the code
# is correct and runs fine; this comment just silences the false positive.

import os
from datetime import datetime

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle,
    NextPageTemplate, PageTemplate, Frame, BaseDocTemplate
)
from reportlab.pdfgen import canvas

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
GENERATED_DIR = os.path.join(BASE_DIR, "reports", "generated")

CHAPTER_FIELDS = [
    ("Chapter 1: Introduction", ["introduction", "problem_statement", "objectives"]),
    ("Chapter 2: System Analysis", ["existing_system", "proposed_system", "methodology"]),
    ("Chapter 3: System Design & Implementation", ["modules", "technologies"]),
]

FIELD_LABELS = {
    "introduction": "1.1 Introduction",
    "problem_statement": "1.2 Problem Statement",
    "objectives": "1.3 Objectives",
    "existing_system": "2.1 Existing System",
    "proposed_system": "2.2 Proposed System",
    "methodology": "2.3 Methodology",
    "modules": "3.1 Modules",
    "technologies": "3.2 Technologies Used",
}


def _safe(data, key, default="Not specified."):
    val = (data.get(key) or "").strip()
    return val if val else default


# ---------------------------------------------------------------------------
# PDF GENERATION
# ---------------------------------------------------------------------------

def _pdf_header_footer(canvas_obj, doc, data):
    canvas_obj.saveState()
    width, height = A4
    # Header
    canvas_obj.setFont("Helvetica-Bold", 9)
    canvas_obj.setFillColor(colors.HexColor("#4b2e83"))
    canvas_obj.drawString(0.75 * inch, height - 0.55 * inch, data.get("college", "College Name"))
    canvas_obj.setFont("Helvetica", 8)
    canvas_obj.setFillColor(colors.HexColor("#666666"))
    canvas_obj.drawRightString(width - 0.75 * inch, height - 0.55 * inch,
                                data.get("project_title", "Project Report")[:60])
    canvas_obj.setStrokeColor(colors.HexColor("#4b2e83"))
    canvas_obj.line(0.75 * inch, height - 0.62 * inch, width - 0.75 * inch, height - 0.62 * inch)

    # Footer
    canvas_obj.setFont("Helvetica", 8)
    canvas_obj.setFillColor(colors.HexColor("#666666"))
    canvas_obj.line(0.75 * inch, 0.65 * inch, width - 0.75 * inch, 0.65 * inch)
    canvas_obj.drawString(0.75 * inch, 0.45 * inch, "Project Report Generator (Template-Based)")
    canvas_obj.drawRightString(width - 0.75 * inch, 0.45 * inch, f"Page {doc.page}")
    canvas_obj.restoreState()


def generate_pdf(data, output_path):
    """
    data: dict with keys matching the Create Report form fields.
    output_path: absolute path where the PDF should be written.
    """
    doc = BaseDocTemplate(
        output_path, pagesize=A4,
        leftMargin=0.9 * inch, rightMargin=0.9 * inch,
        topMargin=0.9 * inch, bottomMargin=0.9 * inch,
        title=data.get("project_title", "Project Report"),
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin,
                   doc.width, doc.height, id="normal")

    def on_page(c, d):
        _pdf_header_footer(c, d, data)

    template = PageTemplate(id="main", frames=[frame], onPage=on_page)
    doc.addPageTemplates([template])

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CoverTitle", fontSize=22, leading=28,
                               alignment=TA_CENTER, textColor=colors.HexColor("#2d1b52"),
                               fontName="Helvetica-Bold", spaceAfter=14))
    styles.add(ParagraphStyle(name="CoverSub", fontSize=13, leading=18,
                               alignment=TA_CENTER, textColor=colors.HexColor("#333333"),
                               fontName="Helvetica"))
    styles.add(ParagraphStyle(name="SectionHeading", fontSize=15, leading=20,
                               textColor=colors.HexColor("#2d1b52"), spaceBefore=14,
                               spaceAfter=8, fontName="Helvetica-Bold"))
    styles.add(ParagraphStyle(name="SubHeading", fontSize=12, leading=16,
                               textColor=colors.HexColor("#4b2e83"), spaceBefore=10,
                               spaceAfter=6, fontName="Helvetica-Bold"))
    styles.add(ParagraphStyle(name="BodyJustify", fontSize=10.5, leading=16,
                               alignment=TA_JUSTIFY, fontName="Helvetica",
                               spaceAfter=8))
    styles.add(ParagraphStyle(name="TOCItem", fontSize=11, leading=20,
                               fontName="Helvetica"))

    story = []

    # ---------------- Cover Page ----------------
    story.append(Spacer(1, 0.6 * inch))
    story.append(Paragraph(data.get("college", "College Name"), styles["CoverSub"]))
    story.append(Spacer(1, 0.15 * inch))
    story.append(Paragraph(f"Department of {data.get('department', 'Computer Science')}",
                            styles["CoverSub"]))
    story.append(Spacer(1, 0.6 * inch))
    story.append(Paragraph("A Project Report on", styles["CoverSub"]))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(data.get("project_title", "Project Title"), styles["CoverTitle"]))
    story.append(Spacer(1, 0.3 * inch))
    story.append(Paragraph(f"Submitted for the {data.get('project_type', 'Mini Project')} "
                            f"in partial fulfillment of the requirements of "
                            f"{data.get('branch', 'B.Tech')}", styles["CoverSub"]))
    story.append(Spacer(1, 0.6 * inch))

    cover_table = Table([
        ["Submitted by:", data.get("student_name", "Student Name")],
        ["Roll Number:", data.get("roll_number", "-")],
        ["Guide:", data.get("guide_name", "-")],
        ["Academic Year:", data.get("academic_year", "-")],
    ], colWidths=[1.8 * inch, 3.2 * inch])
    cover_table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 11),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
    ]))
    story.append(cover_table)
    story.append(PageBreak())

    # ---------------- Certificate ----------------
    story.append(Paragraph("CERTIFICATE", styles["SectionHeading"]))
    story.append(Spacer(1, 0.2 * inch))
    cert_text = (
        f"This is to certify that <b>{data.get('student_name', 'Student Name')}</b> "
        f"(Roll No: {data.get('roll_number', '-')}) of the Department of "
        f"{data.get('department', '-')}, {data.get('college', '-')}, has successfully "
        f"completed the {data.get('project_type', 'Mini Project')} titled "
        f"\"<b>{data.get('project_title', '-')}</b>\" during the academic year "
        f"{data.get('academic_year', '-')} under my/our guidance, in partial "
        f"fulfillment of the requirements for the degree of {data.get('branch', 'B.Tech')}."
    )
    story.append(Paragraph(cert_text, styles["BodyJustify"]))
    story.append(Spacer(1, 0.6 * inch))
    story.append(Paragraph(f"Guide: {data.get('guide_name', '-')}", styles["BodyJustify"]))
    story.append(Paragraph("Head of Department", styles["BodyJustify"]))
    story.append(PageBreak())

    # ---------------- Acknowledgement ----------------
    story.append(Paragraph("ACKNOWLEDGEMENT", styles["SectionHeading"]))
    ack_text = (
        f"I would like to express my sincere gratitude to my guide, "
        f"{data.get('guide_name', 'my guide')}, for the continuous support, "
        f"guidance and encouragement throughout this {data.get('project_type', 'project')}. "
        f"I am also thankful to the Department of {data.get('department', '-')} and "
        f"{data.get('college', 'my college')} for providing the resources and "
        f"opportunity to carry out this work."
    )
    story.append(Paragraph(ack_text, styles["BodyJustify"]))
    story.append(Spacer(1, 0.4 * inch))
    story.append(Paragraph(data.get("student_name", "Student Name"), styles["BodyJustify"]))
    story.append(PageBreak())

    # ---------------- Abstract ----------------
    story.append(Paragraph("ABSTRACT", styles["SectionHeading"]))
    story.append(Paragraph(_safe(data, "abstract"), styles["BodyJustify"]))
    story.append(PageBreak())

    # ---------------- Table of Contents ----------------
    story.append(Paragraph("TABLE OF CONTENTS", styles["SectionHeading"]))
    toc_items = [
        "Certificate", "Acknowledgement", "Abstract",
        "Chapter 1: Introduction",
        "Chapter 2: System Analysis",
        "Chapter 3: System Design & Implementation",
        "Advantages & Limitations",
        "Future Scope", "Conclusion", "References", "Appendix",
    ]
    for i, item in enumerate(toc_items, 1):
        story.append(Paragraph(f"{i}. {item}", styles["TOCItem"]))
    story.append(PageBreak())

    # ---------------- Chapters ----------------
    for chap_title, fields in CHAPTER_FIELDS:
        story.append(Paragraph(chap_title, styles["SectionHeading"]))
        for f in fields:
            story.append(Paragraph(FIELD_LABELS.get(f, f.title()), styles["SubHeading"]))
            story.append(Paragraph(_safe(data, f), styles["BodyJustify"]))
        story.append(PageBreak())

    # ---------------- Advantages / Limitations ----------------
    story.append(Paragraph("Advantages & Limitations", styles["SectionHeading"]))
    story.append(Paragraph("Advantages", styles["SubHeading"]))
    story.append(Paragraph(_safe(data, "advantages"), styles["BodyJustify"]))
    story.append(Paragraph("Limitations", styles["SubHeading"]))
    story.append(Paragraph(_safe(data, "limitations"), styles["BodyJustify"]))
    story.append(PageBreak())

    # ---------------- Future Scope ----------------
    story.append(Paragraph("Future Scope", styles["SectionHeading"]))
    story.append(Paragraph(_safe(data, "future_scope"), styles["BodyJustify"]))
    story.append(PageBreak())

    # ---------------- Conclusion ----------------
    story.append(Paragraph("Conclusion", styles["SectionHeading"]))
    story.append(Paragraph(_safe(data, "conclusion"), styles["BodyJustify"]))
    story.append(PageBreak())

    # ---------------- References ----------------
    story.append(Paragraph("References", styles["SectionHeading"]))
    refs = _safe(data, "references", "1. No references provided.")
    for i, line in enumerate(refs.split("\n"), 1):
        line = line.strip()
        if not line:
            continue
        if not line[0].isdigit():
            line = f"{i}. {line}"
        story.append(Paragraph(line, styles["BodyJustify"]))
    story.append(PageBreak())

    # ---------------- Appendix ----------------
    story.append(Paragraph("Appendix", styles["SectionHeading"]))
    story.append(Paragraph(
        "Supporting screenshots, code snippets and additional material referenced "
        "in this report can be attached here.", styles["BodyJustify"]))

    doc.build(story)
    return output_path


# ---------------------------------------------------------------------------
# DOCX GENERATION
# ---------------------------------------------------------------------------

def _set_cell_text(cell, text, bold=False, size=11):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def _add_page_number_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def generate_docx(data, output_path):
    document = Document()

    # Base style
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    section = document.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Header
    header_p = section.header.paragraphs[0]
    header_p.text = f"{data.get('college', 'College Name')}"
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Footer with page number
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.add_run("Project Report Generator (Template-Based)  |  Page ")
    _add_page_number_field(footer_p)

    def heading(text, level=1, color="2D1B52"):
        h = document.add_heading(level=level)
        run = h.add_run(text)
        run.font.color.rgb = RGBColor.from_string(color)
        return h

    def body(text):
        p = document.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p

    # ---------------- Cover ----------------
    for _ in range(3):
        document.add_paragraph()
    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(data.get("college", "College Name"))
    run.font.size = Pt(16)
    run.bold = True

    dept_p = document.add_paragraph()
    dept_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dept_p.add_run(f"Department of {data.get('department', 'Computer Science')}").italic = True

    document.add_paragraph()
    lead_p = document.add_paragraph()
    lead_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead_p.add_run("A Project Report on")

    proj_title_p = document.add_paragraph()
    proj_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = proj_title_p.add_run(data.get("project_title", "Project Title"))
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("2D1B52")

    sub_p = document.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.add_run(f"Submitted for the {data.get('project_type', 'Mini Project')} in partial "
                  f"fulfillment of the requirements of {data.get('branch', 'B.Tech')}")

    document.add_paragraph()
    table = document.add_table(rows=4, cols=2)
    table.autofit = True
    rows_data = [
        ("Submitted by:", data.get("student_name", "Student Name")),
        ("Roll Number:", data.get("roll_number", "-")),
        ("Guide:", data.get("guide_name", "-")),
        ("Academic Year:", data.get("academic_year", "-")),
    ]
    for i, (label, value) in enumerate(rows_data):
        _set_cell_text(table.cell(i, 0), label, bold=True)
        _set_cell_text(table.cell(i, 1), value)

    document.add_page_break()

    # ---------------- Certificate ----------------
    heading("CERTIFICATE", level=1)
    body(
        f"This is to certify that {data.get('student_name', 'Student Name')} "
        f"(Roll No: {data.get('roll_number', '-')}) of the Department of "
        f"{data.get('department', '-')}, {data.get('college', '-')}, has successfully "
        f"completed the {data.get('project_type', 'Mini Project')} titled "
        f"\"{data.get('project_title', '-')}\" during the academic year "
        f"{data.get('academic_year', '-')} under my/our guidance, in partial fulfillment "
        f"of the requirements for the degree of {data.get('branch', 'B.Tech')}."
    )
    document.add_paragraph()
    body(f"Guide: {data.get('guide_name', '-')}")
    body("Head of Department")
    document.add_page_break()

    # ---------------- Acknowledgement ----------------
    heading("ACKNOWLEDGEMENT", level=1)
    body(
        f"I would like to express my sincere gratitude to my guide, "
        f"{data.get('guide_name', 'my guide')}, for the continuous support, guidance "
        f"and encouragement throughout this {data.get('project_type', 'project')}. "
        f"I am also thankful to the Department of {data.get('department', '-')} and "
        f"{data.get('college', 'my college')} for providing the resources and "
        f"opportunity to carry out this work."
    )
    document.add_paragraph()
    body(data.get("student_name", "Student Name"))
    document.add_page_break()

    # ---------------- Abstract ----------------
    heading("ABSTRACT", level=1)
    body(_safe(data, "abstract"))
    document.add_page_break()

    # ---------------- Table of Contents ----------------
    heading("TABLE OF CONTENTS", level=1)
    toc_items = [
        "Certificate", "Acknowledgement", "Abstract",
        "Chapter 1: Introduction",
        "Chapter 2: System Analysis",
        "Chapter 3: System Design & Implementation",
        "Advantages & Limitations",
        "Future Scope", "Conclusion", "References", "Appendix",
    ]
    for i, item in enumerate(toc_items, 1):
        document.add_paragraph(f"{i}. {item}")
    document.add_page_break()

    # ---------------- Chapters ----------------
    for chap_title, fields in CHAPTER_FIELDS:
        heading(chap_title, level=1)
        for f in fields:
            heading(FIELD_LABELS.get(f, f.title()), level=2, color="4B2E83")
            body(_safe(data, f))
        document.add_page_break()

    # ---------------- Advantages / Limitations ----------------
    heading("Advantages & Limitations", level=1)
    heading("Advantages", level=2, color="4B2E83")
    body(_safe(data, "advantages"))
    heading("Limitations", level=2, color="4B2E83")
    body(_safe(data, "limitations"))
    document.add_page_break()

    # ---------------- Future Scope ----------------
    heading("Future Scope", level=1)
    body(_safe(data, "future_scope"))
    document.add_page_break()

    # ---------------- Conclusion ----------------
    heading("Conclusion", level=1)
    body(_safe(data, "conclusion"))
    document.add_page_break()

    # ---------------- References ----------------
    heading("References", level=1)
    refs = _safe(data, "references", "1. No references provided.")
    for i, line in enumerate(refs.split("\n"), 1):
        line = line.strip()
        if not line:
            continue
        if not line[0].isdigit():
            line = f"{i}. {line}"
        document.add_paragraph(line)
    document.add_page_break()

    # ---------------- Appendix ----------------
    heading("Appendix", level=1)
    body("Supporting screenshots, code snippets and additional material referenced "
         "in this report can be attached here.")

    document.save(output_path)
    return output_path


def generate_reports(data, base_filename):
    """Generate both PDF and DOCX, return (pdf_path, docx_path) absolute paths."""
    os.makedirs(GENERATED_DIR, exist_ok=True)
    pdf_path = os.path.join(GENERATED_DIR, f"{base_filename}.pdf")
    docx_path = os.path.join(GENERATED_DIR, f"{base_filename}.docx")
    generate_pdf(data, pdf_path)
    generate_docx(data, docx_path)
    return pdf_path, docx_path